from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\Documents\ChatGPT\Thailand Disaster Watch v1.0")
OUT = ROOT / "output" / "manual"
OUT.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT / "Thailand-Disaster-Watch-Easy-User-Guide-v1.3-Expanded.docx"

ATTACH = Path(r"C:\Users\User\AppData\Local\Temp")
IMAGES = {
    "dashboard": ATTACH / "codex-clipboard-d80dec9c-b280-49f7-849e-1a940be5283c.png",
    "about": ATTACH / "codex-clipboard-366d99af-f18a-4706-b575-24e9d5c3cfd6.png",
    "province": ATTACH / "codex-clipboard-17e3f9eb-0ed7-4600-99bd-61594c2d1f65.png",
    "seismo": ATTACH / "codex-clipboard-eb21fbc0-23c9-4ba6-91b9-1ea92c5e04c6.png",
    "mysites": ATTACH / "codex-clipboard-a23ed114-8bcb-4633-be81-e6ceff9b50b5.png",
    "bcm": ATTACH / "codex-clipboard-f843420c-ce48-4ec5-8d4f-1985e1d614b3.png",
    "chao_phraya_flow": ATTACH / "codex-clipboard-b9e87b28-a65d-48b6-a063-515cac659400.png",
    "windy_wind": ATTACH / "codex-clipboard-52a4156f-236a-459e-b7bb-56683df0f133.png",
    "windy_rain": ATTACH / "codex-clipboard-1748497f-8c3d-4a0b-82df-918e63b64b7e.png",
    "windy_clouds": ATTACH / "codex-clipboard-95d101fb-1369-4731-9c8c-a2105e714267.png",
    "windy_storm": ATTACH / "codex-clipboard-b14bf2d6-62de-452f-a496-44bc2440cfcf.png",
    "windy_cams": ATTACH / "codex-clipboard-7eae5fc4-b7af-4e64-9f8c-0a08055fd4bf.png",
}

for key, path in IMAGES.items():
    if not path.exists():
        raise FileNotFoundError(f"Missing image {key}: {path}")


NAVY = "0B172A"
BLUE = "0EA5E9"
CYAN = "38BDF8"
GREEN = "10B981"
PALE_BLUE = "EAF6FC"
PALE_GREEN = "EAF8F3"
PALE_AMBER = "FFF6E5"
PALE_RED = "FDECEC"
MID_GRAY = "5B6472"
LIGHT_GRAY = "E5EAF0"
WHITE = "FFFFFF"
BLACK = "111827"
FONT = "Sarabun"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths_in: list[float], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(round(w * 1440) for w in widths_in)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_in:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(col)

    for row in table.rows:
        for idx, (cell, width) in enumerate(zip(row.cells, widths_in)):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, bold=False, color=BLACK, italic=False) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text: str, *, size=11, bold=False, color=BLACK, align=None, before=0, after=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    set_run(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text: str, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    if level == 1:
        set_run(r, size=16, bold=True, color="075985")
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        set_run(r, size=13, bold=True, color="0369A1")
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
    else:
        set_run(r, size=12, bold=True, color="0C4A6E")
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
    return p


def add_bullet(doc, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), bold=True)
        set_run(p.add_run(text[len(bold_prefix):]))
    else:
        set_run(p.add_run(text))
    return p


def add_step(doc, number: int, title: str, detail: str):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [0.55, 5.95])
    left, right = table.rows[0].cells
    set_cell_shading(left, BLUE)
    set_cell_shading(right, PALE_BLUE)
    left_p = left.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(left_p.add_run(str(number)), size=15, bold=True, color=WHITE)
    right_p = right.paragraphs[0]
    right_p.paragraph_format.space_after = Pt(2)
    set_run(right_p.add_run(title), size=11.5, bold=True, color="075985")
    detail_p = right.add_paragraph()
    detail_p.paragraph_format.space_after = Pt(0)
    detail_p.paragraph_format.line_spacing = 1.2
    set_run(detail_p.add_run(detail), size=10.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_callout(doc, label: str, text: str, fill: str, accent: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(label + "  "), size=11, bold=True, color=accent)
    set_run(p.add_run(text), size=10.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def set_alt_text(inline_shape, title: str, description: str):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_image(doc, image_path: Path, caption: str, alt: str, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(image_path), width=Inches(width))
    set_alt_text(shape, caption, alt)
    cp = doc.add_paragraph(style="Caption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(3)
    cp.paragraph_format.space_after = Pt(8)
    cp.paragraph_format.keep_with_next = False
    set_run(cp.add_run(caption), size=9.5, color=MID_GRAY, italic=True)


def add_page_break(doc):
    doc.add_page_break()


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run(run, size=9, color=MID_GRAY)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for style_name in ("Heading 1", "Heading 2", "Heading 3", "Caption", "List Bullet", "List Number"):
    st = styles[style_name]
    st.font.name = FONT
    st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

# Running header/footer
header_p = section.header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(header_p.add_run("THAILAND DISASTER WATCH  |  EASY USER GUIDE"), size=8.5, bold=True, color="64748B")
footer_p = section.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(footer_p.add_run("FutureGreen Disaster Intelligence Platform  •  หน้า "), size=9, color=MID_GRAY)
add_page_field(footer_p)

# Cover page - editorial_cover pattern with compact-reference preset and Sarabun override.
add_text(doc, "THAILAND DISASTER WATCH", size=11, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
logo = ROOT / "public" / "futuregreen-logo.png"
if logo.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(logo), width=Inches(1.0))
    set_alt_text(shape, "FutureGreen logo", "ตราสัญลักษณ์ FutureGreen Disaster Intelligence Platform")
add_text(doc, "คู่มือการใช้งานแบบง่าย", size=28, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=4)
add_text(doc, "รู้จักหน้าจอหลัก เลือกพื้นที่ และใช้ข้อมูลประกอบการติดตามสถานการณ์", size=14, color="334155", align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
add_callout(
    doc,
    "สำหรับผู้ใช้ครั้งแรก",
    "ใช้เวลาประมาณ 10 นาที คู่มือนี้อธิบายจากหน้าจอจริงของระบบเวอร์ชัน 1.3.0",
    PALE_BLUE,
    "0369A1",
)
add_image(
    doc,
    IMAGES["dashboard"],
    "ภาพรวมหน้า National Situation Monitoring",
    "หน้าหลักของ Thailand Disaster Watch แสดงแผนที่ประเทศไทย เมนูทางซ้าย และข้อมูลสถานการณ์ทางขวา",
    width=6.25,
)
add_text(doc, "Production URL: https://disaster.futuregreennet.com", size=10.5, bold=True, color="0369A1", align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
add_text(doc, "ฉบับย่อสำหรับผู้ใช้งานทั่วไป  |  Version 1.3.0", size=9.5, color=MID_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

add_page_break(doc)

add_heading(doc, "1. ระบบนี้ใช้ทำอะไร", 1)
add_text(doc, "Thailand Disaster Watch เป็นเว็บแอปสำหรับรวมข้อมูลเชิงพื้นที่และข้อมูลสถานการณ์ไว้ในหน้าจอเดียว เพื่อช่วยให้ผู้ใช้เห็นภาพรวมก่อนเปิดดูรายละเอียดจากแหล่งข้อมูลต้นทาง")

features = [
    ("ดูภาพรวมระดับประเทศ", "เปิดแผนที่ประเทศไทยและเลือกจังหวัดหรือภูมิภาคที่ต้องการติดตาม"),
    ("ติดตามสภาพอากาศและฝน", "เปิดหน้า Weather, เรดาร์ และ Windy เพื่อดูสภาพอากาศประกอบ"),
    ("ดูสถานการณ์น้ำ", "เข้าถึงข้อมูลแม่น้ำ เขื่อน ผังน้ำเจ้าพระยา และลิงก์แหล่งข้อมูลที่เกี่ยวข้อง"),
    ("เฝ้าระวังแผ่นดินไหว", "ใช้ SeismoWatch ดูเหตุการณ์จากเครือข่ายแผ่นดินไหวที่ระบบแสดง"),
    ("ติดตามพื้นที่ธุรกิจ", "ใช้ My Sites เพื่อรวมโรงงาน สำนักงาน คลังสินค้า หรือพื้นที่สำคัญ"),
    ("จัดทำรายงาน BCM", "เปิดรายงานสรุปความเสี่ยงเพื่อใช้ประกอบการประชุมหรือบันทึกเป็น PDF"),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
set_table_widths(table, [1.85, 4.65])
set_repeat_table_header(table.rows[0])
for idx, title in enumerate(("ความสามารถ", "ใช้งานอย่างไร")):
    set_cell_shading(table.rows[0].cells[idx], NAVY)
    p = table.rows[0].cells[idx].paragraphs[0]
    set_run(p.add_run(title), size=10.5, bold=True, color=WHITE)
for title, detail in features:
    cells = table.add_row().cells
    set_run(cells[0].paragraphs[0].add_run(title), size=10.5, bold=True, color="075985")
    set_run(cells[1].paragraphs[0].add_run(detail), size=10.5)
    for cell in cells:
        set_cell_margins(cell)
set_table_widths(table, [1.85, 4.65])

add_heading(doc, "ก่อนเริ่มใช้งาน", 2)
add_step(doc, 1, "เปิดเว็บไซต์", "เข้า https://disaster.futuregreennet.com ผ่าน Chrome, Edge หรือ Safari")
add_step(doc, 2, "ตรวจชื่อพื้นที่", "ดูชื่อประเทศ ภูมิภาค หรือจังหวัดที่แถบด้านบนก่อนอ่านข้อมูล")
add_step(doc, 3, "ตรวจเวลาและแหล่งข้อมูล", "ดู Observed, Retrieved, Freshness และ Attribution ในแผงข้อมูลด้านขวา")
add_step(doc, 4, "เปิดแหล่งข้อมูลต้นทางเมื่อจำเป็น", "หากใช้เพื่อการตัดสินใจสำคัญ ให้กดลิงก์หน่วยงานทางการเพื่อตรวจสอบซ้ำ")

add_callout(
    doc,
    "ข้อควรระวัง",
    "ระบบนี้เป็น Decision Support ไม่ใช่ประกาศเตือนภัยทางการ และไม่ควรใช้เป็นแหล่งเดียวในการตัดสินใจฉุกเฉิน",
    PALE_RED,
    "B91C1C",
)

add_page_break(doc)

add_heading(doc, "2. รู้จักหน้าจอหลัก", 1)
add_image(
    doc,
    IMAGES["dashboard"],
    "หน้าจอหลัก: เมนูคำสั่ง แผนที่ และ Situation Overview",
    "ภาพหน้าจอหลักแสดงเมนูด้านซ้าย แผนที่ประเทศไทยตรงกลาง และแผง Situation Overview ด้านขวา",
)
add_heading(doc, "ส่วนประกอบสำคัญ", 2)
add_bullet(doc, "1) เมนูด้านซ้าย: เลือก GIS Map, Weather, Windy, SeismoWatch, ผังน้ำเจ้าพระยา, My Sites และ About")
add_bullet(doc, "2) ช่องค้นหา: พิมพ์ชื่อจังหวัดหรือ Region เพื่อไปยังพื้นที่ได้เร็ว")
add_bullet(doc, "3) แผนที่กลาง: ซูม เลื่อน และเลือกจังหวัดเพื่อดูสถานการณ์เฉพาะพื้นที่")
add_bullet(doc, "4) แถบสถานะใต้แผนที่: บอกว่าหมวด Flood, Rain, River, Dam, Alerts และ CCTV พร้อมหรือไม่มีข้อมูล")
add_bullet(doc, "5) Situation Overview: แสดงผู้ให้ข้อมูล เวลาเรียกข้อมูล ความสด และลิงก์ตรวจสอบต้นทาง")

add_callout(
    doc,
    "อ่านสถานะให้ถูกต้อง",
    "No live data หรือ UNKNOWN หมายถึงยังไม่มีข้อมูลที่ยืนยันได้ ไม่ได้หมายความว่าสถานการณ์ปกติ",
    PALE_AMBER,
    "92400E",
)

add_page_break(doc)

add_heading(doc, "3. เลือกจังหวัดและดูสถานการณ์พื้นที่", 1)
add_image(
    doc,
    IMAGES["province"],
    "ตัวอย่างหน้า Pathum Thani Situation",
    "หน้าสถานการณ์จังหวัดปทุมธานี แสดงขอบเขตจังหวัดบนแผนที่และข้อมูลประจำพื้นที่ด้านขวา",
)
add_step(doc, 1, "ค้นหาจังหวัด", "พิมพ์ชื่อจังหวัดในช่องค้นหา หรือเลือกจากกลุ่มภูมิภาคทางซ้าย")
add_step(doc, 2, "ยืนยันพื้นที่บนหัวหน้า", "ชื่อหน้าและ Breadcrumb ต้องตรงกับจังหวัดที่ต้องการ")
add_step(doc, 3, "อ่าน Situation Overview", "ตรวจ Source, Provider, Dataset, Observed, Retrieved, Freshness และ Confidence")
add_step(doc, 4, "เปิดข้อมูลเฉพาะเรื่อง", "เลือก Weather, Windy, SeismoWatch หรือข้อมูลน้ำตามวัตถุประสงค์")
add_step(doc, 5, "กลับสู่ประเทศไทย", "กด Reset Thailand หรือเลือกประเทศไทยจาก Navigation")

add_heading(doc, "ความหมายของคำสำคัญ", 2)
terms = [
    ("Observed", "เวลาที่ข้อมูลถูกตรวจวัดหรือสังเกต"),
    ("Published", "รอบการเผยแพร่จากผู้ให้ข้อมูล"),
    ("Retrieved", "เวลาที่ระบบดึงข้อมูลเข้ามา"),
    ("Freshness", "ความสดหรือความล่าช้าของข้อมูล"),
    ("Confidence", "ระดับความพร้อมหรือการตรวจสอบของชุดข้อมูล"),
    ("Attribution", "ชื่อหน่วยงานหรือผู้ให้ข้อมูลที่ต้องอ้างอิง"),
]
term_table = doc.add_table(rows=0, cols=2)
term_table.style = "Table Grid"
for term, meaning in terms:
    cells = term_table.add_row().cells
    set_cell_shading(cells[0], PALE_BLUE)
    set_run(cells[0].paragraphs[0].add_run(term), size=10.5, bold=True, color="075985")
    set_run(cells[1].paragraphs[0].add_run(meaning), size=10.5)
set_table_widths(term_table, [1.45, 5.05])

add_heading(doc, "4. ใช้โมดูลติดตามสถานการณ์", 1)
add_heading(doc, "GIS Map View", 2)
add_text(doc, "ใช้ดูขอบเขตพื้นที่ เลือกจังหวัด และเปิดหรือปิด Layer ที่เกี่ยวข้อง เช่น ขอบเขตจังหวัด เรดาร์ และพื้นที่น้ำท่วมจากดาวเทียม")
add_heading(doc, "Weather และ Windy", 2)
add_text(doc, "ใช้ดูสภาพอากาศ ภาพเรดาร์ ลม เมฆ ฝน และแบบจำลองพายุ ควรตรวจเวลาอัปเดตและประเภทข้อมูลว่าเป็น Observed หรือ Forecast ทุกครั้ง")
add_heading(doc, "ผังน้ำเจ้าพระยา", 2)
add_text(doc, "ใช้เปิดภาพรวมเส้นทางการไหลของน้ำในลุ่มน้ำเจ้าพระยาและลิงก์ไปยังข้อมูล HII/ThaiWater เพื่อดูรายละเอียดเพิ่มเติม")

add_page_break(doc)

add_heading(doc, "4.1 ผังน้ำลุ่มเจ้าพระยา", 1)
add_image(
    doc,
    IMAGES["chao_phraya_flow"],
    "หน้าผังน้ำลุ่มเจ้าพระยาและแบบจำลองเส้นทางน้ำ",
    "หน้าผังน้ำลุ่มเจ้าพระยา แสดงลิงก์ HII Portal และ ThaiWater จุดควบคุมสำคัญ และผังเส้นทางน้ำพร้อมระยะเวลาเดินทางโดยประมาณ",
    width=6.25,
)
add_heading(doc, "ใช้หน้านี้อย่างไร", 2)
add_bullet(doc, "กด เปิดผังล่าสุด (HII.or.th) เมื่อต้องการตรวจผังหรือข้อมูลจากระบบต้นทาง")
add_bullet(doc, "กด รายงานเขื่อนใหญ่ (ThaiWater) เพื่อตรวจปริมาณน้ำกักเก็บและข้อมูลเขื่อนจากหน้าต้นทาง")
add_bullet(doc, "ใช้รายการเกณฑ์ควบคุมด้านซ้ายเพื่อรู้จักจุดเฝ้าระวังสำคัญ เช่น สถานี C.2 และเขื่อนเจ้าพระยา")
add_bullet(doc, "ใช้ผังด้านขวาเพื่อทำความเข้าใจลำดับเส้นทางน้ำและเวลาเดินทางโดยประมาณ ไม่ใช่เวลาถึงที่ยืนยันแล้ว")
add_callout(
    doc,
    "ข้อควรระวังเรื่องเวลาเดินทางน้ำ",
    "ระยะเวลาในผังเป็นข้อมูลประกอบการวางแผน อาจเปลี่ยนตามการระบายน้ำ ปริมาณฝน สภาพลำน้ำ และการบริหารเขื่อน ต้องตรวจข้อมูล HII, ThaiWater และประกาศหน่วยงานทางการก่อนตัดสินใจ",
    PALE_AMBER,
    "92400E",
)

add_page_break(doc)

add_heading(doc, "4.2 Windy: ดูลม ฝน เมฆ พายุ และภาพพื้นที่", 1)
add_image(
    doc,
    IMAGES["windy_wind"],
    "โหมด Wind แสดงทิศทางและความเร็วลม",
    "หน้า Windy โหมด Wind แสดงแผนที่ประเทศไทย เส้นทางลม สเกลความเร็วลม ปุ่มเลือก Layer และแถบเวลา",
    width=6.25,
)
add_heading(doc, "ส่วนควบคุมที่ใช้บ่อย", 2)
add_bullet(doc, "เลือก Layer ด้านบน: Wind, Rain, Clouds, Storm Tracker, Live Cams, Pressure หรือ Temp")
add_bullet(doc, "คลิกตำแหน่งบนแผนที่เพื่ออ่านค่าของจุดนั้น พร้อมตรวจหน่วยที่แสดง เช่น km/h หรือ mm")
add_bullet(doc, "เลื่อนแถบเวลาด้านล่างเพื่อดูช่วงเวลาที่ต้องการ และตรวจวัน/เวลาทุกครั้งก่อนบันทึกภาพ")
add_bullet(doc, "กด เปิดใน Windy.com เมื่อต้องการดูรายละเอียด โมเดล หรือเครื่องมือเพิ่มเติมจากต้นทาง")
add_callout(
    doc,
    "Windy เป็นข้อมูลแบบจำลอง",
    "สี เส้นลม และค่าบนแผนที่เปลี่ยนตาม Layer เวลา และโมเดลที่เลือก จึงต้องระบุ Layer เวลา หน่วย และโมเดลเมื่อใช้อ้างอิง และไม่ใช้แทนประกาศเตือนภัยทางการ",
    PALE_AMBER,
    "92400E",
)

add_page_break(doc)

add_heading(doc, "Windy: ฝนและเมฆ", 2)
add_image(
    doc,
    IMAGES["windy_rain"],
    "โหมด Rain ใช้ดูพื้นที่และช่วงเวลาที่แบบจำลองแสดงฝน",
    "หน้า Windy โหมด Rain แสดงแถบสีปริมาณฝน หน่วยมิลลิเมตร เส้นลม แผนที่ และแถบเวลา",
    width=5.70,
)
add_bullet(doc, "Rain: อ่านพื้นที่สีร่วมกับสเกลหน่วย mm และเวลาที่เลือก ค่าที่เห็นเป็นภาพตาม Layer/โมเดล ไม่ใช่เครื่องวัดฝนหน้างาน")
add_image(
    doc,
    IMAGES["windy_clouds"],
    "โหมด Clouds ใช้ดูเมฆและรายละเอียดของจุดที่เลือก",
    "หน้า Windy โหมด Clouds แสดงชั้นเมฆ เส้นลม ค่าร้อยละเมฆและค่าประกอบของจุดที่คลิก",
    width=5.70,
)
add_bullet(doc, "Clouds: ใช้ดูรูปแบบและความหนาแน่นของเมฆจาก Layer ที่เลือก ควรอ่านร่วมกับ Rain และข้อมูลเรดาร์เมื่อมี")

add_page_break(doc)

add_heading(doc, "Windy: ติดตามพายุและกล้องเว็บแคม", 2)
add_image(
    doc,
    IMAGES["windy_storm"],
    "โหมด Storm Tracker มีปุ่มเปิดระบบติดตามพายุจาก Windy",
    "หน้า Windy โหมด Storm Tracker แสดงแถบแจ้งว่าต้องเปิดระบบติดตามพายุเต็มจอจาก Windy และยังคงเห็นแผนที่ลมประกอบ",
    width=5.70,
)
add_bullet(doc, "Storm Tracker: ใช้สำรวจแนวโน้มและเส้นทางพายุเบื้องต้น แล้วตรวจยืนยันกับ TMD, DDPM หรือประกาศทางการ")
add_image(
    doc,
    IMAGES["windy_cams"],
    "โหมด Live Cams มีปุ่มเปิดจุดกล้องเว็บแคมจาก Windy",
    "หน้า Windy โหมด Live Cams แสดงแถบเปิดจุดกล้องเว็บแคมและแผนที่ลมประกอบ",
    width=5.70,
)
add_bullet(doc, "Live Cams: ใช้ดูสภาพพื้นที่ด้วยภาพประกอบเท่านั้น กล้องอาจออฟไลน์ ภาพล่าช้า มุมมองจำกัด หรือไม่ครอบคลุมพื้นที่เสี่ยง")
add_callout(
    doc,
    "ยืนยันก่อนสั่งการ",
    "ข้อมูลจาก Windy และกล้องเว็บแคมช่วยให้เห็นภาพรวม แต่ไม่ใช่หลักฐานเดียวสำหรับการหยุดผลิต เปลี่ยนเส้นทางขนส่ง หรือสั่งอพยพ",
    PALE_RED,
    "B91C1C",
)

add_page_break(doc)

add_heading(doc, "SeismoWatch", 2)
add_image(
    doc,
    IMAGES["seismo"],
    "SeismoWatch: แผนที่ติดตามแผ่นดินไหวและสึนามิ",
    "หน้าจอ SeismoWatch แสดงจำนวนเหตุการณ์ แผนที่โลก ตัวกรองขนาดแผ่นดินไหว และรายการเหตุการณ์ล่าสุด",
)
add_bullet(doc, "เลือกช่วงเวลา เช่น 24 Hours เพื่อจำกัดเหตุการณ์ที่ต้องการดู")
add_bullet(doc, "ใช้ตัวกรอง M4+, M5+ หรือ M6+ เพื่อดูเหตุการณ์ตาม Magnitude")
add_bullet(doc, "เลือก TH Near Thailand เพื่อเน้นเหตุการณ์ใกล้ประเทศไทย")
add_bullet(doc, "ตรวจ Source และเวลาเหตุการณ์ก่อนนำข้อมูลไปใช้อ้างอิง")
add_callout(doc, "เหตุการณ์แผ่นดินไหว", "จุดบนแผนที่คือข้อมูลเหตุการณ์ที่ระบบรวบรวม ไม่ใช่คำสั่งอพยพ หากมีความเสี่ยงให้ตรวจประกาศ TMD, DDPM, NOAA/PTWC หรือหน่วยงานทางการ", PALE_AMBER, "92400E")

add_page_break(doc)

add_heading(doc, "5. My Sites: ติดตามโรงงาน คลังสินค้า และพื้นที่สำคัญ", 1)
add_image(
    doc,
    IMAGES["mysites"],
    "My Sites: รวมพื้นที่ธุรกิจที่ต้องเฝ้าระวัง",
    "หน้าจอ My Sites แสดงรายการโรงงาน สำนักงาน คลังสินค้า และนิคมอุตสาหกรรม พร้อมปุ่มดูแผนที่ สภาพอากาศ และรายงาน BCM",
)
add_text(doc, "My Sites ช่วยให้ผู้ใช้รวมพื้นที่ที่สำคัญต่อธุรกิจไว้ในหน้าจอเดียว เหมาะสำหรับฝ่าย Production, Logistics, Facility, Safety และ BCM")
add_heading(doc, "ปุ่มที่ใช้บ่อย", 2)
add_bullet(doc, "ดูแผนที่ GIS: เปิดตำแหน่งของ Site และบริเวณรอบข้าง")
add_bullet(doc, "ตรวจสภาพอากาศ: เปิดข้อมูลสภาพอากาศของจังหวัดหรือพิกัดที่เกี่ยวข้อง")
add_bullet(doc, "สรุปรายงาน BCM: เปิดรายงานประเมินความเสี่ยงเพื่อทบทวนและพิมพ์")
add_callout(
    doc,
    "การรักษาความลับ",
    "ชื่อบริษัท ที่อยู่ พิกัด และผลประเมิน BCM อาจเป็นข้อมูลธุรกิจภายใน ควรกำหนดสิทธิ์ผู้เข้าถึงและตรวจสอบก่อนแชร์หรือเผยแพร่คู่มือ",
    PALE_RED,
    "B91C1C",
)

add_page_break(doc)

add_heading(doc, "6. เปิดและบันทึกรายงาน BCM", 1)
add_image(
    doc,
    IMAGES["bcm"],
    "ตัวอย่างรายงานประเมินความเสี่ยงและความต่อเนื่องทางธุรกิจ",
    "หน้าต่าง BCM Report แสดงสถานะรวม ปัจจัยฝน น้ำท่วม การระบายน้ำ ผลกระทบจากเขื่อน และมาตรการแนะนำ",
    width=5.55,
)
add_step(doc, 1, "เลือก Site", "เปิด My Sites แล้วเลือกรายการโรงงาน คลังสินค้า หรือพื้นที่ที่ต้องการ")
add_step(doc, 2, "กดสรุปรายงาน BCM", "ระบบจะแสดงหน้าต่างรายงานพร้อมชื่อพื้นที่และเวลาประเมิน")
add_step(doc, 3, "ทบทวนข้อมูลและคำแนะนำ", "ตรวจฝน น้ำท่วมจากดาวเทียม การระบายน้ำ ผลกระทบจากเขื่อน และมาตรการที่แนะนำ")
add_step(doc, 4, "พิมพ์หรือบันทึก PDF", "กด Print BCM Report แล้วเลือก Save as PDF ในหน้าต่างพิมพ์")
add_callout(
    doc,
    "ก่อนอนุมัติการดำเนินงาน",
    "รายงานต้องแสดงข้อมูลที่ตรวจสอบได้และเวลาอัปเดตจริง หากข้อมูลไม่พร้อม ให้ถือสถานะเป็น UNKNOWN และตรวจหน่วยงานต้นทางก่อน",
    PALE_AMBER,
    "92400E",
)

add_heading(doc, "7. หน้า About และข้อจำกัดของระบบ", 1)
add_image(
    doc,
    IMAGES["about"],
    "หน้า About แสดงวัตถุประสงค์ ผู้พัฒนา และสถานะระบบ",
    "หน้าเกี่ยวกับระบบแสดงคำอธิบายแพลตฟอร์ม คำเตือนสถานะ Development Preview ผู้พัฒนา และเวอร์ชัน",
)
add_text(doc, "หน้า About ใช้ตรวจชื่อระบบ ผู้พัฒนา เวอร์ชัน วัตถุประสงค์ และข้อสงวนสิทธิ์ด้านข้อมูล ผู้ใช้ควรอ่านส่วนนี้ก่อนใช้ระบบในงานที่มีผลต่อความปลอดภัยหรือการดำเนินธุรกิจ")
add_heading(doc, "สิ่งที่ระบบทำได้", 2)
add_bullet(doc, "รวมข้อมูลและลิงก์หลายแหล่งไว้ในหน้าจอเดียว")
add_bullet(doc, "ช่วยให้ค้นหาและเปรียบเทียบพื้นที่ได้เร็วขึ้น")
add_bullet(doc, "ช่วยเตรียมข้อมูลประกอบการประชุมและ BCM")
add_heading(doc, "สิ่งที่ระบบไม่ควรใช้แทน", 2)
add_bullet(doc, "ประกาศเตือนภัยอย่างเป็นทางการ")
add_bullet(doc, "คำสั่งอพยพหรือคำสั่งของหน่วยงานรัฐ")
add_bullet(doc, "การตรวจวัดหน้างานและการยืนยันจากผู้รับผิดชอบพื้นที่")
add_bullet(doc, "การอนุมัติหยุดผลิต ปิดโรงงาน หรือเปลี่ยนแผนขนส่งโดยอัตโนมัติ")

add_page_break(doc)

add_heading(doc, "8. วิธีอ่านข้อมูลอย่างปลอดภัย", 1)
checks = [
    ("พื้นที่ถูกต้องหรือไม่", "ชื่อจังหวัด พิกัด และขอบเขตต้องตรงกับพื้นที่จริง"),
    ("ข้อมูลชนิดใด", "แยก Observed, Forecast, Satellite, Model และ Warning"),
    ("เวลาใด", "ตรวจ Observed, Published และ Retrieved ห้ามใช้วันที่หน้าเว็บเพียงอย่างเดียว"),
    ("แหล่งใด", "ต้องมี Source, Provider และ Attribution ที่ตรวจสอบได้"),
    ("สถานะใด", "AVAILABLE ไม่เท่ากับ NORMAL และ No live data ไม่เท่ากับปลอดภัย"),
    ("ยืนยันซ้ำแล้วหรือไม่", "เปิดลิงก์หน่วยงานต้นทางก่อนตัดสินใจสำคัญ"),
]
for i, (title, detail) in enumerate(checks, 1):
    add_step(doc, i, title, detail)

add_heading(doc, "ตัวอย่างการใช้งานในโรงงาน", 2)
add_bullet(doc, "ฝ่าย Production: ดูสภาพอากาศและความเสี่ยงก่อนประชุม Daily Operation")
add_bullet(doc, "ฝ่าย Logistics: ตรวจเส้นทางและพื้นที่คลังสินค้าก่อนจัดส่ง")
add_bullet(doc, "ฝ่าย Facility: ตรวจพื้นที่ระบายน้ำ ปั๊มน้ำ และเครื่องกำเนิดไฟฟ้า")
add_bullet(doc, "ฝ่าย Safety/BCM: บันทึกหลักฐานและตรวจประกาศต้นทางก่อนยกระดับสถานการณ์")

add_callout(doc, "หลักการสั้นที่สุด", "ข้อมูลไม่ครบ = แสดง UNKNOWN, ข้อมูลเก่า = แสดง STALE, ข้อมูลตัวอย่าง = แสดง DEMO และเหตุฉุกเฉิน = ตรวจประกาศทางการเสมอ", PALE_GREEN, "047857")

add_page_break(doc)

add_heading(doc, "9. คำถามที่พบบ่อย", 1)
faq = [
    ("เว็บเข้าไม่ได้ทำอย่างไร", "ตรวจอินเทอร์เน็ต ลอง Refresh หนึ่งครั้ง และตรวจหน้า Status/ช่องทางผู้ดูแลระบบ"),
    ("แผนที่ไม่ขึ้นทำอย่างไร", "รอสักครู่ ตรวจการเชื่อมต่อ และลองเปลี่ยน Base Map หากยังไม่ขึ้นให้บันทึกภาพ Error ส่งผู้ดูแล"),
    ("ทำไมบางหมวดขึ้น No live data", "หมายถึงระบบยังไม่มีข้อมูลสดที่ใช้แสดงในขณะนั้น ไม่ได้หมายความว่าไม่มีภัย"),
    ("สามารถใช้แทนประกาศ TMD หรือ DDPM ได้หรือไม่", "ไม่ได้ ต้องตรวจประกาศจากหน่วยงานทางการทุกครั้ง"),
    ("สามารถแชร์ BCM Report ได้หรือไม่", "แชร์ได้เฉพาะผู้มีสิทธิ์และต้องตรวจว่ารายงานไม่มีข้อมูลลับหรือพิกัดที่ห้ามเปิดเผย"),
    ("ข้อมูลผิดหรือเก่าควรทำอย่างไร", "หยุดใช้ข้อมูลรายการนั้น บันทึกเวลาและภาพหน้าจอ แล้วแจ้งผู้ดูแลระบบพร้อมลิงก์หน้าเว็บ"),
]
for q, a in faq:
    add_heading(doc, q, 2)
    add_text(doc, a, size=10.5, after=5)

add_heading(doc, "ติดต่อและข้อมูลอ้างอิง", 1)
add_text(doc, "เว็บไซต์: https://disaster.futuregreennet.com", bold=True, color="0369A1")
add_text(doc, "Repository: https://github.com/TheSor55/thailand-disaster-watch", color="0369A1")
add_text(doc, "ผู้สร้างโครงการและ Lead Developer: Sorawit Suwannarong", bold=True)
add_text(doc, "เอกสารฉบับนี้จัดทำจากหน้าจอระบบ Version 1.3.0 ณ วันที่ 27 สิงหาคม 2569", size=9.5, color=MID_GRAY)
add_callout(doc, "หมายเหตุ", "หน้าจอและชื่อเมนูอาจเปลี่ยนแปลงเมื่อระบบอัปเดต ควรตรวจ Version บนหน้า About ให้ตรงกับคู่มือ", PALE_BLUE, "0369A1")

doc.core_properties.title = "คู่มือการใช้งาน Thailand Disaster Watch แบบง่าย"
doc.core_properties.subject = "คู่มือผู้ใช้ระบบติดตามสถานการณ์ภัยพิบัติและ BCM"
doc.core_properties.author = "FutureGreen Disaster Intelligence Platform"
doc.core_properties.keywords = "Thailand Disaster Watch, GIS, BCM, User Guide"
doc.core_properties.comments = "Generated as an easy-start user guide for version 1.3.0"

doc.save(DOCX_PATH)
print(DOCX_PATH)
